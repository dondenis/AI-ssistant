import os
import sys
import re
from datetime import datetime
from flask import Flask, request, render_template_string, send_file, redirect, url_for
from werkzeug.utils import secure_filename
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill, Font
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document

# Ensure the gpt4free subfolder is in sys.path
current_dir = os.path.dirname(os.path.abspath(__file__))
gpt4free_path = os.path.join(current_dir, 'gpt4free')
if gpt4free_path not in sys.path:
    sys.path.insert(0, gpt4free_path)
from g4f.client import Client
client = Client()  # default client

app = Flask(__name__)

# --- Configuration ---
UPLOAD_FOLDER = "inputs"
OUTPUT_FOLDER = "outputs"
ALLOWED_EXTENSIONS = {'docx'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# --- Helper Functions ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_docx(file_path):
    """
    Reads a DOCX file and returns a list of its paragraph texts.
    """
    doc = Document(file_path)
    return [para.text for para in doc.paragraphs]

# --- Transcript Processing with 3 API Calls on the Entire Transcript ---

def process_entire_transcript(transcript, interviewer_name, progress_callback=None):
    """
    Processes the entire transcript through three API calls:
    
    1. Grammar Correction:
       "Please correct the grammar and spelling for the following transcript while preserving its structure.
        Each line starts with a timestamp (HH:MM:SS), followed by the speaker's name and text."
    2. Quote Extraction:
       "Please extract only the relevant quotes (lines) spoken by the interviewee (ignore any line where the speaker's name is '{interviewer_name}').
        If a line contains multiple key concepts, split them into separate lines. Output each quote on a new line in the format:
        'HH:MM:SS Speaker: Quote'."
    3. Final Refinement:
       "Please refine each of the following quotes to be concise and memorable, and at the end of each quote append in brackets the categorization 
        into one of these themes: Business Model, Market Outlook, or Challenges. Maintain the format 'HH:MM:SS Speaker: Quote [Theme]'."
    
    The function prints intermediate API outputs for debugging and invokes the progress_callback(progress) function (if provided)
    with progress values: 25, 50, 75, 100.
    
    Returns the final refined transcript as a text block.
    """
    # API Call 1: Grammar Correction
    if progress_callback: progress_callback(25)
    prompt1 = (
        "Please correct the grammar and spelling for the following transcript while preserving its structure. "
        "Each line starts with a timestamp (HH:MM:SS), followed by the speaker's name and then the text. "
        "Output the corrected transcript in the same format.\n\n" + transcript
    )
    response1 = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt1}],
        web_search=False
    )
    corrected_transcript = response1.choices[0].message.content.strip()
    print("DEBUG: Corrected Transcript:")
    print(corrected_transcript)
    if progress_callback: progress_callback(50)
    
    # API Call 2: Extract Relevant Quotes (only interviewee lines)
    prompt2 = (
        f"Please extract only the relevant quotes from the following transcript that are spoken by the interviewee. "
        f"Ignore any lines where the speaker's name is '{interviewer_name}'. If a line contains multiple key concepts, split them into separate quotes. "
        "Output each quote on a new line in the format: 'HH:MM:SS Speaker: Quote'.\n\n" + corrected_transcript
    )
    response2 = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt2}],
        web_search=False
    )
    extracted_quotes = response2.choices[0].message.content.strip()
    print("DEBUG: Extracted Quotes:")
    print(extracted_quotes)
    if progress_callback: progress_callback(75)
    
    # API Call 3: Final Refinement of Quotes with Categorization
    prompt3 = (
        "Please refine each of the following quotes to be concise, memorable, and well-formatted. "
        "For each quote, append at the end in brackets the categorization (one of: Business Model, Market Outlook, Challenges). "
        "Maintain the format 'HH:MM:SS Speaker: Quote [Theme]'.\n\n" + extracted_quotes
    )
    response3 = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt3}],
        web_search=False
    )
    final_quotes = response3.choices[0].message.content.strip()
    print("DEBUG: Final Refined Quotes:")
    print(final_quotes)
    if progress_callback: progress_callback(100)
    
    return final_quotes

def parse_quotes(final_quotes, interview_filename):
    """
    Parses the final refined transcript into a list of segments.
    Assumes each line is in the format:
       HH:MM:SS Speaker: Quote text [Theme]
    Extracts timestamp, speaker, quote text, and the topic from the bracket.
    Returns a list of dictionaries with keys:
       interview_filename, timestamp, topic, quote.
    """
    segments = []
    # Regex to capture timestamp, speaker, quote, and optional theme in brackets
    pattern = re.compile(r"^(\d{2}:\d{2}:\d{2})\s+([^:]+):\s*(.*?)(?:\s+\[(.*?)\])?\s*$", re.MULTILINE)
    for match in pattern.finditer(final_quotes):
        timestamp, speaker, quote, theme = match.groups()
        # Use the theme if provided, otherwise default to "Interviewee"
        topic = theme if theme in ["Business Model", "Market Outlook", "Challenges"] else "Uncategorized"
        segments.append({
            "interview_filename": interview_filename,
            "timestamp": timestamp,
            "topic": topic,
            "quote": quote  # the quote without the trailing bracketed topic
        })
    return segments

def append_to_excel(all_segments, excel_path):
    """
    Creates an Excel file with columns:
      Interview File Name | Timestamp | Topic | Quote
    Applies formatting:
      - Hides gridlines
      - Auto-adjusts column widths
      - Header row highlighted in light green and bold
      - Entire data formatted as an Excel table (filterable)
    """
    wb = Workbook()
    ws = wb.active
    ws.append(["Interview File Name", "Timestamp", "Topic", "Quote"])
    
    for seg in all_segments:
        ws.append([seg["interview_filename"], seg["timestamp"], seg["topic"], seg["quote"]])
    
    ws.sheet_view.showGridLines = False
    
    for col in ws.columns:
        max_length = 0
        column_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = max_length + 2

    header_fill = PatternFill("solid", fgColor="CCFFCC")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True)
    
    last_row = ws.max_row
    if last_row > 1:
        table_ref = f"A1:D{last_row}"
        table = Table(displayName="InterviewTable", ref=table_ref)
        style = TableStyleInfo(name="TableStyleMedium2", showRowStripes=False, showColumnStripes=False)
        table.tableStyleInfo = style
        ws.add_table(table)
    
    wb.save(excel_path)
    return excel_path

# --- Flask Web Interface ---
INDEX_HTML = """
<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8">
    <title>Interview Processor</title>
    <style>
      /* Use a darker beige for background */
      body {
        background-color: #e9e4db;
        font-family: Arial, sans-serif;
        display: flex;
        flex-direction: column;
        align-items: center;
        margin: 0;
        padding: 20px;
      }
      h2, h3 {
        text-align: center;
      }
      /* White for drop zone box */
      #drop_zone {
        width: 300px;
        height: 150px;
        border: 2px dashed #ccc;
        border-radius: 10px;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        text-align: center;
        margin: 20px;
        background-color: #fff;
        cursor: pointer;
      }
      #drop_zone.hover {
        border-color: #333;
      }
      #file_input {
        display: none;
      }
      /* White for list container */
      #file_list_container {
        margin-top: 15px;
        width: 300px;
        background-color: #fff;
        border: 1px solid #ccc;
        border-radius: 5px;
        padding: 10px;
      }
      /* Orange button, slightly grayed out if disabled */
      .button {
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        margin-top: 10px;
        cursor: pointer;
      }
      .button:disabled {
        background-color: #ccc;
        cursor: not-allowed;
      }
      /* Dark orange active button */
      .button-active {
        background-color: #D87A56;
        color: #fff;
      }
      #progressBarContainer {
        width: 300px;
        background-color: #ddd;
        border-radius: 5px;
        overflow: hidden;
        margin-top: 10px;
        display: none;
      }
      #progressBar {
        width: 0%;
        height: 20px;
        background-color: #76c7c0;
        text-align: center;
        color: #fff;
        line-height: 20px;
      }
      .example {
        border: 1px solid #ccc;
        padding: 10px;
        margin-top: 20px;
        background-color: #fff;
        width: 80%;
        max-width: 600px;
      }
      table {
        width: 100%;
        border-collapse: collapse;
      }
      table, th, td {
        border: 1px solid #ddd;
      }
      th, td {
        padding: 8px;
        text-align: center;
      }
      th {
        background-color: #f7e1a0;
      }
      #interviewer_input {
        margin-top: 10px;
        padding: 8px;
        width: 280px;
        border: 1px solid #ccc;
        border-radius: 5px;
      }
    </style>
  </head>
  <body>
    <h2>Interview Processor</h2>
    <div id="drop_zone">Drag and drop your DOCX file<br/>or click to select</div>
    <input type="file" id="file_input" accept=".docx" multiple>
    <div id="file_list_container" style="display:none;">
      <h3>Currently Included Interviews</h3>
      <ul id="file_list"></ul>
      <button class="button" id="reset_btn">Reset Queue</button>
    </div>
    <input type="text" id="interviewer_input" placeholder="Enter interviewer name (to ignore their quotes)">
    <button class="button" id="generate_btn" disabled>Generate Excel</button>
    <div id="progressBarContainer">
      <div id="progressBar">0%</div>
    </div>
    <div class="example">
      <h3>Example</h3>
      <p><strong>Input Transcript (sample):</strong></p>
      <p>00:01:23 John: Our company’s business model focuses on subscription revenue and innovative strategies.</p>
      <p>00:01:45 Mary: Yes, we have a diverse portfolio.</p>
      <p><strong>Expected Excel Output:</strong></p>
      <table>
        <tr>
          <th>Interview File Name</th>
          <th>Timestamp</th>
          <th>Topic</th>
          <th>Quote</th>
        </tr>
        <tr>
          <td>sample_interview.docx</td>
          <td>00:01:23</td>
          <td>Business Model</td>
          <td>Our company’s business model focuses on subscription revenue and innovative strategies.</td>
        </tr>
      </table>
    </div>
    <script>
      const dropZone = document.getElementById('drop_zone');
      const fileInput = document.getElementById('file_input');
      const fileListContainer = document.getElementById('file_list_container');
      const fileList = document.getElementById('file_list');
      const generateBtn = document.getElementById('generate_btn');
      const resetBtn = document.getElementById('reset_btn');
      const progressBarContainer = document.getElementById('progressBarContainer');
      const progressBar = document.getElementById('progressBar');
      const interviewerInput = document.getElementById('interviewer_input');

      let uploadedFiles = [];

      dropZone.addEventListener('dragover', (e) => {
        e.preventDefault();
        dropZone.classList.add('hover');
      });

      dropZone.addEventListener('dragleave', () => {
        dropZone.classList.remove('hover');
      });

      dropZone.addEventListener('drop', (e) => {
        e.preventDefault();
        dropZone.classList.remove('hover');
        if(e.dataTransfer.files.length) {
          addFiles(e.dataTransfer.files);
        }
      });

      dropZone.addEventListener('click', () => {
        fileInput.click();
      });

      fileInput.addEventListener('change', () => {
        if(fileInput.files.length) {
          addFiles(fileInput.files);
        }
      });

      function addFiles(fileListObj) {
        for(let i = 0; i < fileListObj.length; i++){
          if(fileListObj[i].name.toLowerCase().endsWith('.docx')){
            uploadedFiles.push(fileListObj[i]);
          }
        }
        displayFileList();
      }

      function displayFileList(){
        if(uploadedFiles.length > 0){
          fileListContainer.style.display = 'block';
          let html = "";
          uploadedFiles.forEach((f) => {
            html += "<li>" + f.name + "</li>";
          });
          fileList.innerHTML = html;
          // Enable the button
          generateBtn.disabled = false;
          generateBtn.classList.add('button-active');
        } else {
          fileListContainer.style.display = 'none';
          generateBtn.disabled = true;
          generateBtn.classList.remove('button-active');
        }
      }

      resetBtn.addEventListener('click', () => {
        uploadedFiles = [];
        fileList.innerHTML = "";
        fileListContainer.style.display = 'none';
        generateBtn.disabled = true;
        generateBtn.classList.remove('button-active');
      });

        generateBtn.addEventListener('click', () => {
    if(uploadedFiles.length === 0) return;
    progressBarContainer.style.display = 'block';
    updateProgress(33); // First touchpoint: when the first API call is sent
    uploadFiles();
  });

  async function uploadFiles(){
    try {
      const formData = new FormData();
      for(let i = 0; i < uploadedFiles.length; i++){
        formData.append('files', uploadedFiles[i]);
      }
      formData.append('interviewer', interviewerInput.value);

      const response = await fetch("/generate_excel", {
        method: "POST",
        body: formData
      });
      updateProgress(66); // Second touchpoint: when the last API call is completed
      if(!response.ok) {
        alert("Failed to export file: " + response.statusText);
        return;
      }
      const blob = await response.blob();
      updateProgress(100); // Third touchpoint: when the Excel is ready to download
      const url = window.URL.createObjectURL(blob);
      const downloadName = "interview_analysis.xlsx";
      const a = document.createElement("a");
      a.href = url;
      a.download = downloadName;
      document.body.appendChild(a);
      a.click();
      a.remove();
    } catch(err) {
      alert("Failed to export file: " + err);
    }
  }

  function updateProgress(value){
    progressBar.style.width = value + '%';
    progressBar.textContent = value + '%';
  }

      async function uploadAllFiles(){
        try {
          const formData = new FormData();
          for(let i = 0; i < uploadedFiles.length; i++){
            formData.append('files', uploadedFiles[i]);
          }
          formData.append('interviewer', interviewerInput.value);

          const response = await fetch("/generate_excel", {
            method: "POST",
            body: formData
          });
          if(!response.ok) {
            alert("Failed to export file: " + response.statusText);
            return;
          }
          const blob = await response.blob();
          const url = window.URL.createObjectURL(blob);
          // Always name the final Excel "interview_analysis.xlsx"
          const downloadName = "interview_analysis.xlsx";
          const a = document.createElement("a");
          a.href = url;
          a.download = downloadName;
          document.body.appendChild(a);
          a.click();
          a.remove();
        } catch(err) {
          alert("Failed to export file: " + err);
        }
      }
    </script>
  </body>
</html>
"""

@app.route("/", methods=["GET"])
def index():
    return render_template_string(INDEX_HTML)

@app.route("/generate_excel", methods=["POST"])
def generate_excel():
    files = request.files.getlist("files")
    interviewer = request.form.get("interviewer", "").strip()
    if not files:
        return "No files provided", 400

    all_segments = []
    for file in files:
        if file.filename == "" or not allowed_file(file.filename):
            continue
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        paragraphs = read_docx(file_path)
        transcript = "\n".join(paragraphs)
        # Process the entire transcript with the three API calls.
        final_quotes_text = process_entire_transcript(transcript, interviewer, progress_callback=lambda p: None)
        # Parse the final refined quotes into segments.
        segments = parse_quotes(final_quotes_text, filename)
        all_segments.extend(segments)
    
    if not all_segments:
        return "No valid segments found in the provided files.", 400

    excel_filename = "merged_output.xlsx"
    excel_path = os.path.join(OUTPUT_FOLDER, excel_filename)
    append_to_excel(all_segments, excel_path)
    
    return send_file(excel_path, as_attachment=True)

# def parse_quotes(final_quotes, interview_filename):
#     """
#     Parses the final refined quotes (each on a new line) into a list of segments.
#     Each line should be in the format: 
#       HH:MM:SS Speaker: Quote text [Theme]
#     Extracts the timestamp, speaker, quote, and theme.
#     Returns a list of dictionaries with keys:
#       interview_filename, timestamp, topic, quote.
#     The topic is extracted from the bracketed theme and then removed from the quote.
#     """
#     segments = []
#     pattern = re.compile(r"^(\d{2}:\d{2}:\d{2})\s+([^:]+):\s*(.*?)(?:\s+\[(.*?)\])?\s*$", re.MULTILINE)
#     for match in pattern.finditer(final_quotes):
#         timestamp, speaker, quote, theme = match.groups()
#         topic = theme if theme in ["Business Model", "Market Outlook", "Challenges"] else "Uncategorized"
#         segments.append({
#             "interview_filename": interview_filename,
#             "timestamp": timestamp,
#             "topic": topic,
#             "quote": quote
#         })
#     return segments

if __name__ == "__main__":
    app.run(debug=True)
